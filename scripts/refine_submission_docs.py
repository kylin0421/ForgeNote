"""Rebuild the competition development and test manuals.

The generated documents use one shared compact-reference style so the
submission package reads as a coherent set. All technical claims are tied to
the current repository implementation and executed tests.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "to_be_submitted_docs"
OUTPUT_DIR = Path(os.environ.get("SUBMISSION_DOCS_OUTPUT_DIR", SUBMISSION))
ASSETS = ROOT / "docs" / "assets"
EVIDENCE = SUBMISSION / "assets" / "acceptance-test-report.png"
VIDEO_EVIDENCE = SUBMISSION / "assets" / "explainer-video-test-report.png"

NAVY = "0B2545"
BLUE = "2563EB"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
MUTED = "52647A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "087F5B"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111827"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120


def _rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def _set_run_font(
    run,
    *,
    latin: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if italic is not None:
        run.italic = italic


def _set_style_font(
    style, *, size: float, color: str = BLACK, bold: bool = False
) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = _rgb(color)
    style.font.bold = bold


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _table_borders(table, color: str = "C9D4E2", size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    _set_run_font(run, size=9, color=MUTED)


def _add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def _configure_document(doc: Document, running_label: str) -> dict[str, int]:
    output_settings = doc.settings._element
    if output_settings.find(qn("w:doNotAutoCompressPictures")) is None:
        output_settings.append(OxmlElement("w:doNotAutoCompressPictures"))

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, HEADING_BLUE, 18, 10),
        ("Heading 2", 13, HEADING_BLUE, 14, 7),
        ("Heading 3", 12, HEADING_DARK, 10, 5),
    ):
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles = doc.styles
    for name in (
        "DocTitle",
        "DocSubtitle",
        "Kicker",
        "Lead",
        "Caption",
        "CodeBlock",
        "Small",
    ):
        if name not in styles:
            styles.add_style(name, 1)

    _set_style_font(styles["DocTitle"], size=28, color=NAVY, bold=True)
    styles["DocTitle"].paragraph_format.space_after = Pt(8)
    styles["DocTitle"].paragraph_format.keep_with_next = True
    _set_style_font(styles["DocSubtitle"], size=13.5, color=MUTED)
    styles["DocSubtitle"].paragraph_format.space_after = Pt(18)
    _set_style_font(styles["Kicker"], size=10, color=BLUE, bold=True)
    styles["Kicker"].paragraph_format.space_after = Pt(4)
    _set_style_font(styles["Lead"], size=11.5, color=NAVY, bold=True)
    styles["Lead"].paragraph_format.space_after = Pt(7)
    styles["Lead"].paragraph_format.line_spacing = 1.25
    _set_style_font(styles["Caption"], size=9, color=MUTED)
    styles["Caption"].font.italic = True
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].paragraph_format.space_before = Pt(4)
    styles["Caption"].paragraph_format.space_after = Pt(8)
    _set_style_font(styles["CodeBlock"], size=8.4, color=BLACK)
    styles["CodeBlock"].font.name = "Consolas"
    styles["CodeBlock"]._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    styles["CodeBlock"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    styles["CodeBlock"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["CodeBlock"].paragraph_format.space_after = Pt(0)
    styles["CodeBlock"].paragraph_format.line_spacing = 1.0
    _set_style_font(styles["Small"], size=9, color=MUTED)
    styles["Small"].paragraph_format.space_after = Pt(4)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(running_label)
    _set_run_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Forgenote · ")
    _set_run_font(r, size=9, color=MUTED)
    _add_page_number(p)

    return {
        "bullet": _add_numbering_definition(doc, "bullet"),
        "decimal": _add_numbering_definition(doc, "decimal"),
    }


def _add_cover(
    doc: Document,
    *,
    title: str,
    subtitle: str,
    version: str,
    document_type: str,
    takeaway: str,
) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(58)
    p = doc.add_paragraph("中国软件杯 A3 · 参赛配套文档", style="Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(title, style="DocTitle")
    doc.add_paragraph(subtitle, style="DocSubtitle")

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目", "Forgenote（智学工坊）个性化学习多智能体系统"),
        ("文档类型", document_type),
        ("版本", version),
        ("编制日期", "2026 年 7 月 19 日"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
            if cell is row.cells[0]:
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, size=9.5, bold=True, color=NAVY)
            else:
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, size=9.5, color=BLACK)
    _set_table_geometry(table, [1900, 7460])
    _table_borders(table)

    doc.add_paragraph()
    _add_callout(doc, "核心结论", takeaway, fill=CALLOUT, color=NAVY)
    doc.add_page_break()


def _add_callout(
    doc: Document, label: str, text: str, *, fill: str, color: str
) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}｜")
    _set_run_font(r, size=10.5, bold=True, color=color)
    r = p.add_run(text)
    _set_run_font(r, size=10.5, color=BLACK)
    _set_table_geometry(table, [CONTENT_WIDTH_DXA])
    _table_borders(table, color="D6E0EC", size=5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def _add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        _set_run_font(first, bold=True, color=NAVY)
        rest = p.add_run(text[len(bold_lead) :])
        _set_run_font(rest)
    else:
        run = p.add_run(text)
        _set_run_font(run)


def _add_bullets(doc: Document, items: Iterable[str], num_id: int) -> None:
    # Word may continue numbering across sections when a single numId is reused.
    # Give each logical list its own definition so bullets remain bullets.
    num_id = _add_numbering_definition(doc, "bullet")
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        _apply_num(p, num_id)
        run = p.add_run(item)
        _set_run_font(run)


def _add_steps(doc: Document, items: Iterable[str], num_id: int) -> None:
    # A fresh definition also guarantees every numbered procedure restarts at 1.
    num_id = _add_numbering_definition(doc, "decimal")
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        _apply_num(p, num_id)
        run = p.add_run(item)
        _set_run_font(run)


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
    *,
    header_fill: str = LIGHT_BLUE,
    small: bool = False,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    _repeat_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        cell.text = text
        _set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            _set_run_font(run, size=9 if small else 9.5, bold=True, color=NAVY)

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = str(value)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if index == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _set_run_font(run, size=8.6 if small else 9.2, color=BLACK)
        if len(table.rows) % 2 == 1:
            for cell in row.cells:
                _set_cell_shading(cell, "FAFCFE")
    _set_table_geometry(table, widths_dxa)
    _table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_image(doc: Document, path: Path, caption: str, *, width: float = 6.25) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def _add_code_listing(doc: Document, source: Path) -> None:
    text = source.read_text(encoding="utf-8").splitlines()
    for line_no, line in enumerate(text, start=1):
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.keep_together = False
        p.paragraph_format.keep_with_next = False
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F7F9FC")
        p._p.get_or_add_pPr().append(shd)
        run = p.add_run(f"{line_no:03d}  {line}")
        _set_run_font(
            run, latin="Consolas", east_asia="Microsoft YaHei", size=8.4, color=BLACK
        )


def build_development_manual() -> Path:
    doc = Document()
    nums = _configure_document(doc, "Forgenote · A3 SYSTEM DEVELOPMENT MANUAL")
    _add_cover(
        doc,
        title="Forgenote 系统开发说明书",
        subtitle="基于大模型的个性化资源生成与学习多智能体系统",
        version="V2.2（带字幕讲解视频版）",
        document_type="系统开发说明书",
        takeaway=(
            "项目的核心创新是由 9 个职责清晰、过程可观测的智能体共同维护学习闭环，并将文本、结构化题库、"
            "图谱、代码、图片、播客和带字幕讲解视频统一纳入来源约束的生成体系。学习画像、课程资料、练习反馈和"
            "路径调整在同一学习记录中持续积累，学生可以核验来源、观察过程并修正画像。"
        ),
    )

    _add_heading(doc, "文档摘要", 1)
    _add_body(
        doc,
        "Forgenote 面向高校课程自主学习场景，基于 Open Notebook 的资料、笔记与 RAG 底座进行赛题化二次开发。"
        "系统把被动的资料问答升级为主动学习工作台：学生通过自然语言表达背景与目标，系统维护动态画像，"
        "由多角色智能体协同完成资料搜集、资源生成、练习、路径规划、即时辅导、效果评估与安全检查。",
    )
    _add_body(
        doc,
        "本说明书以当前仓库代码和 2026 年 7 月 19 日实测结果为依据。重点展开 5 项创新：画像即来源（Profile-as-Source）、"
        "9 角色模式路由、来源约束的多模态资产流水线、基于真实语音时间轴的带字幕讲解视频生成，以及长任务可观测与防幻觉阻断。",
    )
    _add_heading(doc, "赛题要求与实现对照", 2)
    _add_table(
        doc,
        ["官方核心要求", "Forgenote 落地", "可核验实现"],
        [
            [
                "对话式画像 ≥6 维",
                "6 个显式维度 + 稳定画像 + 最近学习信号",
                "LearningProfileDimension；画像事件接口",
            ],
            [
                "多智能体架构",
                "9 个职责角色，按 chat / collect / generate 模式路由",
                "AGENT_BLUEPRINTS；SSE stage trace",
            ],
            [
                "至少 5 类资源",
                "7 类 Studio 资产 + 播客音频 + 讲解视频",
                "requested_outputs；PodcastService；FFmpeg",
            ],
            [
                "动态学习路径",
                "4 阶段路径，含目标、活动、资源与检查点",
                "LearningPathStep；评估与下一步调整",
            ],
            [
                "防幻觉与内容安全",
                "来源优先、空来源硬阻断、结构校验、安全报告",
                "source context；safety_report",
            ],
            [
                "进度追踪",
                "SSE 智能体阶段 + command job 状态/日志",
                "/orchestrate/stream；任务浮窗",
            ],
        ],
        [2200, 3900, 3260],
        small=True,
    )

    doc.add_page_break()
    _add_heading(doc, "1. 需求理解与系统定位", 1)
    _add_heading(doc, "1.1 高校自主学习的核心矛盾", 2)
    _add_table(
        doc,
        ["学生痛点", "直接影响", "系统响应"],
        [
            [
                "资料散乱、质量不一",
                "检索成本高，无法判断先学什么",
                "统一学习记录；资源搜集智能体评分、去重并保留采纳权",
            ],
            [
                "基础与目标差异大",
                "统一讲解与练习无法适配",
                "自然语言画像；生成内容使用专业、目标、历史与偏好",
            ],
            [
                "生成内容难核验",
                "大模型可能脱离课程资料",
                "优先使用已接受来源；引用可见；来源正文缺失时不生成",
            ],
            [
                "长任务等待不透明",
                "学生误以为系统卡死",
                "SSE 显示智能体阶段；后台任务浮窗显示队列、日志与结果",
            ],
        ],
        [2200, 2800, 4360],
    )
    _add_heading(doc, "1.2 产品定位", 2)
    _add_callout(
        doc,
        "定位",
        "面向高校课程自主学习的主动学习系统，以学生确认的课程资料为事实边界，连接画像、资源、练习、路径和反馈。",
        fill=LIGHT_BLUE,
        color=BLUE,
    )
    _add_bullets(
        doc,
        [
            "输入侧：课程资料、网页、文本、笔记、播客字幕、学生目标与学习历史。",
            "理解侧：RAG、语义索引、学习画像和资源质量判断。",
            "行动侧：讲解、测验、闪卡、思维导图、拓展阅读、代码实验、图片、播客与讲解视频。",
            "反馈侧：错题本、学习曲线、画像更新、路径调整和安全提示。",
        ],
        nums["bullet"],
    )

    _add_heading(doc, "2. 创新设计", 1)
    _add_heading(doc, "2.1 创新一：画像即来源（Profile-as-Source）", 2)
    _add_body(
        doc,
        "Forgenote 将“学习画像”保存为学习记录内可查看、可编辑、可检索的特殊 Source。画像与课程资料共享 notebook 上下文，"
        "能够参与内容个性化；学生也可以直接查看证据和置信度，修正不准确的模型判断。",
    )
    _add_table(
        doc,
        ["画像维度", "当前取值来源", "如何随学更新"],
        [
            ["专业背景", "专业、课程与自然语言描述", "首次对话与课程切换"],
            ["知识基础", "学习历史、问答与测验表现", "对话困难点和错题信号"],
            ["学习目标", "goal 字段与请求意图", "chat / collect / generate 请求"],
            ["认知风格", "讲解、练习、代码等组合偏好", "资源使用与生成选择"],
            ["易错点偏好", "错题、混淆表达、失败事件", "quiz_answer / chat_message"],
            ["资源偏好", "采纳来源、上传资料和资产类型", "source_accept 等学习事件"],
        ],
        [1800, 3600, 3960],
        small=True,
    )
    _add_body(
        doc,
        "画像文本区分“稳定画像”和“最近学习信号”。系统最多保留近期 40 条事件，并对目标、易错点和资源偏好做增量归纳，"
        "避免把原始事件无限堆积成不可读日志。",
    )

    _add_heading(doc, "2.2 创新二：9 角色协作与模式路由", 2)
    _add_body(
        doc,
        "9 个智能体分别拥有独立职责、输入输出契约和阶段状态，并由交互模式决定实际参与者。"
        "chat 启用画像、辅导、评估和安全角色；collect 启用画像、课程、搜集和安全角色；generate 执行完整学习链路。",
    )
    _add_callout(
        doc,
        "工程价值",
        "职责隔离让系统可以独立替换模型、缩短轻量交互路径，并把每个阶段作为 SSE 事件呈现给用户。",
        fill=CALLOUT,
        color=NAVY,
    )

    _add_heading(doc, "2.3 创新三：来源约束的多模态资产流水线", 2)
    _add_body(
        doc,
        "学生先选择可信来源，再选择需要的资产类型。文本资产使用学习资产模型生成结构化 JSON；图片资产进入独立图片模型；"
        "播客和讲解视频共享脚本、说话人配置与 TTS 管线。测验、闪卡、思维导图和 Markdown 在保存前执行结构修复与完整性校验。",
    )
    _add_heading(doc, "2.4 创新四：真实语音时间轴驱动的讲解视频", 2)
    _add_body(
        doc,
        "脚本模型在返回逐句播客内容时，同时为概念切换点给出 visual_prompt。TTS 完成后，系统读取每句台词的真实起始时间，"
        "将视觉提示转换为带 time_index 的关键帧计划，再调用图片模型生成 16:9 画面。最后由本地 FFmpeg 按时间轴把图片和播客"
        "音频合成为 H.264/AAC MP4。系统同时依据每段真实语音的起止时间生成 SRT，并把自动换行、分段后的中文字幕烧录到画面。"
        "视频画面、字幕与讲解内容来自同一份脚本，时间对齐可以复现，也不需要视频生成 API。",
    )
    _add_heading(doc, "2.5 创新五：可观测长任务与防幻觉阻断", 2)
    _add_body(
        doc,
        "系统同时提供轻量 SSE 阶段流和可持久化 command job。用户选择的来源尚未完成正文解析时，生成流程返回空资产并提示等待，"
        "从执行层阻止无依据内容进入学习资产；该行为由专项验收用例持续验证。",
    )

    _add_heading(doc, "3. 多智能体架构与实现", 1)
    _add_heading(doc, "3.1 角色分工", 2)
    _add_table(
        doc,
        ["角色", "核心职责", "主要输出"],
        [
            [
                "画像智能体",
                "从自然语言与学习事件中抽取学生状态",
                "6 维画像、置信度、最近信号",
            ],
            ["课程结构智能体", "拆解知识单元并排序先修关系", "课程结构与学习边界"],
            [
                "资源搜集智能体",
                "规划查询、评分、去重、多样性选择",
                "候选来源与采纳状态",
            ],
            ["资源生成智能体", "生成讲解、图谱、阅读等资产", "结构化 LearningResource"],
            [
                "练习实训智能体",
                "生成 Quiz、闪卡、代码实操",
                "题目、答案、解析与实验任务",
            ],
            ["路径规划智能体", "依据画像和掌握度安排顺序", "4 阶段路径与检查点"],
            ["智能辅导智能体", "即时答疑、错误定位、下一步引导", "tutor_answer"],
            ["学习评估智能体", "总结优势、风险与调整建议", "score 与 next_adjustments"],
            [
                "安全校验智能体",
                "检查来源一致性、敏感内容与质量",
                "passed / needs_review 报告",
            ],
        ],
        [1900, 4360, 3100],
        small=True,
    )
    _add_heading(doc, "3.2 编排数据契约", 2)
    _add_body(
        doc,
        "入口对象 LearningOrchestrationRequest 汇总本轮任务所需状态：message、mode、course、major、goal、learning_history、"
        "requested_outputs、accepted_resource_ids、supplemental_materials、learning_record_id、target_language、image_model 以及画像更新开关。",
    )
    _add_body(
        doc,
        "返回对象 LearningOrchestrationResponse 同时携带 profile、collected_resources、resources、learning_path、recommendations、"
        "tutor_answer、evaluation、safety_report 与 trace。前端无需拼接多组不一致数据，就能重放一次完整学习决策。",
    )
    _add_heading(doc, "3.3 执行顺序", 2)
    _add_steps(
        doc,
        [
            "标准化学生请求，补齐课程、专业、目标、历史与目标语言。",
            "读取或创建学习画像来源，并按模式确定本轮活跃角色。",
            "资源搜集智能体规划多个搜索意图，执行检索、启发式评分、可选 LLM 重排与多样性选择。",
            "学生接受来源或上传自有资料；系统构建最多 24,000 字符的来源上下文，单来源最多 7,000 字符。",
            "按 requested_outputs 选择模型与生成管线，校验测验、闪卡、导图与 Markdown 结构。",
            "输出学习路径、辅导、评估和安全报告；将阶段通过 SSE 或后台任务状态展示。",
            "记录本轮学习事件，更新可持久化画像来源，供下一轮学习复用。",
        ],
        nums["decimal"],
    )
    _add_image(
        doc,
        ASSETS / "screenshot-search-agent.png",
        "图 1 资源搜集与学习请求入口：智能体候选来源保留学生确认权",
    )

    _add_heading(doc, "4. 多模态资源生成", 1)
    _add_heading(doc, "4.1 资产类型与呈现", 2)
    _add_table(
        doc,
        ["资产", "生成/处理机制", "主要学习用途"],
        [
            ["课程讲解", "长文 Markdown；章节与自检结构", "建立概念边界与学习顺序"],
            ["Quiz", "结构化题目、选项、答案索引、解析", "学习前后诊断"],
            ["知识闪卡", "正面、背面、提示、证据、来源引用", "高频回忆与易错点复习"],
            ["思维导图", "Mermaid mindmap + 文本/表格降级表示", "结构梳理与关系理解"],
            ["拓展阅读", "按来源和目标组织阅读路径", "扩展深度与权威资料"],
            ["代码实验", "Notebook 任务、运行与误差分析要求", "把概念迁移到实践"],
            ["辅助图片", "独立图片模型生成并保存 PNG/JPEG", "可视化抽象概念"],
            ["播客音频", "提纲 → 对话脚本 → 多说话人 TTS → 合成", "听觉复习与通勤学习"],
            [
                "讲解视频",
                "脚本视觉提示 → 真实 TTS 时间轴 → 图片 + 字幕 → 本地 MP4",
                "图文同步讲解与低成本复习",
            ],
        ],
        [1800, 4200, 3360],
        small=True,
    )
    _add_heading(doc, "4.2 讲解视频生成流程", 2)
    _add_steps(
        doc,
        [
            "脚本模型返回 speaker、dialogue 和可选 visual_prompt；每个分段设置 1–3 个有学习价值的关键画面。",
            "TTS 逐句生成语音并写入真实 start_time；关键帧计划以该时间为准，合并相邻提示并限制最多 12 帧。",
            "图片模型按 16:9 教学画面提示生成图片；每帧保存时间点、台词序号、提示词、模型和供应商信息。",
            "按每段台词的真实 start_time/end_time 生成 SRT；长台词自动换行，并在原语音时段内拆分为两行字幕。",
            "FFmpeg 依据相邻 time_index 计算每张图片的持续时长，将图片、播客音频和中文字幕合成为 1280×720 H.264/AAC MP4。",
            "视频为用户显式勾选的可选输出；图片或合成失败时保留已经完成的播客音频，并返回独立 video_error。",
        ],
        nums["decimal"],
    )
    _add_image(
        doc,
        ASSETS / "screenshot-explainer-video-dialog.png",
        "图 2 高清界面：讲解视频为播客生成任务的可选输出，无需配置视频模型",
        width=6.45,
    )
    _add_heading(doc, "4.3 生成质量控制", 2)
    _add_bullets(
        doc,
        [
            "只返回学生明确选择的资产，避免一次生成过多低价值内容。",
            "学习画像用于调整解释方式，但画像 Source 会从事实证据上下文中剔除，避免把画像元数据生成成课程知识。",
            "闪卡有效数量不足 6 张或缺少证据字段时判定生成无效；Quiz 统一校验答案索引与解析。",
            "Mind map 修复未闭合代码围栏；Markdown 修复破损表格和块边界。",
            "图片模型凭据缺失时给出配置错误，不生成伪图片；播客脚本、图片与 TTS 模型按用途隔离。",
            "讲解视频关键帧、字幕起止时间和字幕分段均来自真实 TTS 时间轴；字幕最多显示两行并直接烧录到成片。",
        ],
        nums["bullet"],
    )
    _add_image(
        doc, ASSETS / "demo-flashcards.jpg", "图 3 真实演示资产：来源约束的知识闪卡"
    )
    _add_image(
        doc, ASSETS / "demo-podcast.jpg", "图 4 真实演示资产：独立播客与 TTS 管线"
    )

    _add_heading(doc, "5. 个性化学习闭环", 1)
    _add_heading(doc, "5.1 四阶段动态路径", 2)
    _add_table(
        doc,
        ["阶段", "目标", "检查点"],
        [
            ["1 建立课程地图", "确认来源并理解主题关系", "能在 3 分钟内讲清知识结构"],
            ["2 补齐关键短板", "Quiz 诊断、错因归类、定向讲解", "基础题正确率达到 80%"],
            [
                "3 完成实操迁移",
                "运行代码、观察参数、解释误差",
                "能解释一次实验结果与误差来源",
            ],
            [
                "4 评估并更新画像",
                "提交反馈、复盘错题、生成下一轮计划",
                "形成下一轮任务清单",
            ],
        ],
        [1900, 4300, 3160],
    )
    _add_heading(doc, "5.2 反馈如何回流", 2)
    _add_body(
        doc,
        "问答困难点、错误答案、资料采纳、资产生成和学习请求会转为学习事件。稳定字段被增量归纳，原始信号保留在最近事件区。"
        "错题本聚合错误题目与解析，学习曲线按日汇总学习量、质量和测验正确率；两者共同为下一轮资源排序提供可解释依据。",
    )
    _add_heading(doc, "5.3 可解释反馈分层", 2)
    _add_table(
        doc,
        ["反馈层", "记录内容", "下一轮用途"],
        [
            [
                "画像层",
                "6 维稳定特征、置信度与证据",
                "决定讲解深度、表达方式与资源偏好",
            ],
            [
                "行为层",
                "问答、资料采纳、资产生成等近期事件",
                "发现短期兴趣、困难与学习节奏变化",
            ],
            [
                "诊断层",
                "错题、解析、正确率与学习曲线",
                "调整路径检查点并生成针对性练习",
            ],
        ],
        [1800, 3800, 3760],
        small=True,
    )
    _add_image(
        doc,
        ASSETS / "screenshot-learning-workspace-hires.png",
        "图 5 高清三栏学习工作台：来源输入、对话辅导与 Studio 资产在同一学习记录中协作",
        width=6.45,
    )

    _add_heading(doc, "6. RAG、防幻觉与内容安全", 1)
    _add_table(
        doc,
        ["控制层", "实现", "作用"],
        [
            ["来源选择", "学生接受候选来源或上传自有资料", "控制事实边界"],
            ["上下文构建", "来源切分、语义索引、token 预算", "减少无关上下文与漂移"],
            ["硬阻断", "来源存在但正文不可用时不生成", "避免模型常识冒充资料结论"],
            ["结构校验", "Quiz、闪卡、导图、Markdown 校验与修复", "阻止不完整资产入库"],
            [
                "安全报告",
                "敏感内容、来源一致性、质量检查",
                "输出 passed 或 needs_review",
            ],
            ["凭据保护", "API key 加密保存且接口不回显明文", "降低凭据泄露风险"],
        ],
        [1700, 4400, 3260],
        small=True,
    )
    _add_callout(
        doc,
        "边界说明",
        "安全机制用于降低而非宣称消除幻觉。真实学术结论仍需学生依据引用来源复核；外部模型服务质量不属于本系统可完全控制的范围。",
        fill="FFF8E8",
        color=GOLD,
    )

    _add_heading(doc, "7. 系统架构与技术实现", 1)
    _add_heading(doc, "7.1 分层架构", 2)
    _add_table(
        doc,
        ["层次", "核心技术", "职责"],
        [
            [
                "交互层",
                "Next.js / React / Radix UI",
                "学习记录、三栏工作台、Studio、错题本、学习曲线",
            ],
            [
                "接口层",
                "FastAPI / SSE",
                "notebook、source、chat、learning、model、credential、command API",
            ],
            [
                "编排层",
                "LearningService / LangGraph 工作流",
                "角色路由、资源搜集、资产生成、路径、评估、安全",
            ],
            [
                "模型适配层",
                "LangChain / OpenAI-compatible / DashScope",
                "文本、Embedding、图片、TTS、STT 多协议统一",
            ],
            [
                "媒体合成层",
                "FFmpeg",
                "依据真实 TTS 时间轴合成画面、音频和字幕；不调用视频生成服务",
            ],
            [
                "数据与任务层",
                "SurrealDB / surreal-commands",
                "来源、资产、画像、模型、凭据、后台任务与日志",
            ],
            [
                "桌面交付层",
                "pywebview / WebView2 / Inno Setup",
                "Windows 安装包、本地服务编排与独立窗口",
            ],
        ],
        [1600, 3300, 4460],
        small=True,
    )
    _add_heading(doc, "7.2 模型用途化配置", 2)
    _add_body(
        doc,
        "系统按用途提供 chat、rag、resource_search、learning_asset、study_guide、quiz、flashcards、"
        "mind_map、reading、code_lab、podcast、embedding、image、text_to_speech、speech_to_text 等用途。"
        "单一资产生成时优先选择对应模型，多资产批量生成时使用 learning_asset 模型；讲解视频复用 podcast、image 和 text_to_speech，"
        "因此无需新增视频模型配置。",
    )
    _add_image(
        doc,
        ASSETS / "screenshot-model-settings.png",
        "图 6 模型/API 配置：基础默认项与按学习用途覆盖",
    )

    _add_heading(doc, "8. 用户界面与体验设计", 1)
    _add_bullets(
        doc,
        [
            "三栏工作台同时展示资料输入、来源感知对话与学习资产，减少页面跳转。",
            "回答支持 Markdown、GFM 表格、公式与引用；学生可框选回答片段继续追问。",
            "右侧 Studio 按课程讲解、测验、闪卡、导图、图片、播客等学习动作组织入口，降低非技术用户配置成本。",
            "长任务通过阶段进度、队列状态、失败日志与结果摘要持续反馈。",
            "错题本与学习曲线位于学习记录顶部独立入口，既突出反馈，又不挤占资料区。",
            "浏览器和 Windows WebView2 共用同一前端，保持演示与实际交付一致。",
        ],
        nums["bullet"],
    )
    _add_heading(doc, "9. 关键接口与数据结构", 1)
    _add_heading(doc, "9.1 学习核心接口", 2)
    _add_table(
        doc,
        ["接口", "作用"],
        [
            [
                "POST /api/learning/orchestrate",
                "同步运行一次学习编排并返回完整闭环对象",
            ],
            [
                "POST /api/learning/orchestrate/stream",
                "通过 SSE 返回各智能体阶段与最终结果",
            ],
            ["POST /api/learning/resource-search/jobs", "提交资源搜集后台任务"],
            ["POST /api/learning/assets/jobs", "按资产类型提交生成任务"],
            ["GET /api/learning/profile-source/{notebook_id}", "获取或创建画像来源"],
            ["POST /api/learning/profile-event", "记录学习事件并按需更新画像"],
            [
                "POST /api/podcasts/generate",
                "提交播客任务，并可通过 generate_video 选择讲解视频",
            ],
            ["GET /api/podcasts/episodes/{id}/video", "读取已合成的讲解视频"],
            ["/api/search 与 /api/context", "全文、向量、语义检索和上下文预算"],
            ["/api/commands", "查询任务状态、日志、结果与失败原因"],
        ],
        [3600, 5760],
        small=True,
    )
    _add_heading(doc, "9.2 核心持久化对象", 2)
    _add_table(
        doc,
        ["对象", "内容"],
        [
            ["notebook", "课程学习记录及归档状态"],
            ["source / source_embedding", "原始资料、全文、主题、分块向量与索引信息"],
            ["note / artifact", "人类笔记与生成学习资产及 notebook 关联"],
            ["model / default models", "模型、供应商、协议和按用途默认配置"],
            ["credential", "加密后的 API key 与各协议端点"],
            ["command job", "后台任务参数、状态、日志、结果和运行时间"],
            [
                "podcast episode",
                "播客脚本、音频、讲解视频、SRT 字幕、关键帧计划与所属 notebook",
            ],
        ],
        [2500, 6860],
        small=True,
    )

    _add_heading(doc, "10. 运行、部署与维护", 1)
    _add_heading(doc, "10.1 三种交付方式", 2)
    _add_table(
        doc,
        ["方式", "适用场景", "入口"],
        [
            [
                "Windows 安装包",
                "答辩与普通用户使用",
                "桌面/开始菜单 Forgenote；安装文件 ForgeNote.exe；WebView2 窗口",
            ],
            ["源码运行", "开发调试", "后端 uv run python run_api.py；前端 npm run dev"],
            ["Docker Compose", "快速复现与服务器部署", "docker compose up -d --build"],
        ],
        [1800, 3300, 4260],
    )
    _add_heading(doc, "10.2 最小运行配置", 2)
    _add_bullets(
        doc,
        [
            "Python 3.11–3.12、Node.js 22+、SurrealDB v2 或 Docker Desktop。",
            "FORGENOTE_ENCRYPTION_KEY、SURREAL_URL、SURREAL_USER、SURREAL_PASSWORD、SURREAL_NAMESPACE=forgenote、SURREAL_DATABASE=forgenote。",
            "至少配置一个文本模型；Embedding、图片、TTS、STT 可按演示需求逐项配置。",
            "Windows 数据默认位于 %LOCALAPPDATA%\\ForgeNote，应用升级不会自动删除用户数据。",
        ],
        nums["bullet"],
    )
    _add_heading(doc, "10.3 维护策略", 2)
    _add_body(
        doc,
        "数据库变更使用 forgenote/database/migrations 管理；新增资产类型需同步扩展请求类型、生成命令、前端 Studio、预览与导出规则。"
        "后台任务故障首先查看 command job 日志；模型故障按分类错误返回可理解提示。",
    )

    _add_heading(doc, "11. 测试与质量证据", 1)
    _add_table(
        doc,
        ["范围", "结果", "说明"],
        [
            ["后端全量回归", "285 passed", "pytest；2 条依赖弃用 warning，无失败"],
            ["学习编排回归", "31 passed", "learning service/API/A3 专项验收"],
            [
                "A3 专项验收",
                "8 passed",
                "画像、9 智能体、7 资产、路径、安全、流式与视频时间轴",
            ],
            [
                "讲解视频专项",
                "11 passed",
                "脚本提示、时间轴、SRT、字幕分段、图片落盘与真实 FFmpeg 合成",
            ],
            [
                "前端单元/组件",
                "10 files / 55 tests passed",
                "Vitest；核心组件与工具函数",
            ],
        ],
        [2300, 2300, 4760],
    )
    _add_body(
        doc,
        "测试结果按实际执行范围分别列出。测试说明书附有 tests/test_submission_acceptance.py、tests/test_explainer_video.py"
        "以及 8/8、11/11 两组通过截图，便于在答辩现场复现。",
    )

    _add_heading(doc, "12. 开源与 AI 工具合规", 1)
    _add_heading(doc, "12.1 开源底座与协议", 2)
    _add_body(
        doc,
        "Forgenote 基于 Open Notebook（https://github.com/lfnovo/open-notebook）二次开发，仓库保留 MIT License 与原版权声明。"
        "团队自主开发集中在学习画像、多智能体学习编排、资源搜索与确认、多类型资产、路径与评估、安全角色、模型用途化配置、"
        "任务可观测、错题本、学习曲线、Windows 交付和比赛材料。",
    )
    _add_heading(doc, "12.2 主要依赖", 2)
    _add_table(
        doc,
        ["类别", "依赖", "用途"],
        [
            ["后端", "FastAPI", "REST API 与 SSE"],
            ["数据库", "SurrealDB", "学习记录、来源、资产、模型和任务"],
            ["AI 编排", "LangGraph / LangChain", "图式工作流和模型调用"],
            ["前端", "Next.js / React", "学习工作台"],
            ["界面", "Radix UI / Tailwind / lucide-react", "组件、布局与图标"],
            ["测试", "pytest / Vitest / Testing Library", "后端、前端与专项验收"],
        ],
        [1700, 3200, 4460],
        small=True,
    )
    _add_heading(doc, "12.3 AI Coding 工具说明", 2)
    _add_body(
        doc,
        "开发过程中使用 AI Coding 工具辅助代码阅读、重构建议、测试执行、错误分析、交互实现与文档整理。"
        "所有关键修改均落入本地仓库，由团队决定取舍，并通过代码审查、自动化测试、lint、build 或手工演示验证；AI 工具不替代最终工程责任。",
    )

    _add_heading(doc, "13. 当前边界与后续规划", 1)
    _add_bullets(
        doc,
        [
            "真实大模型、图片、播客与讲解视频的内容质量受用户模型凭据、额度与网络影响；离线测试验证系统契约和失败处理。",
            "讲解视频采用关键静帧、语音合成和中文字幕，适合知识讲解；当前版本不生成连续角色动作、复杂动画或口型同步。",
            "生产级多租户隔离、长时间稳定性和大规模并发压测仍需在部署环境中补充。",
            "安全报告不能替代学术事实复核；后续可加入引用覆盖率、事实一致性模型和人工审核工作流。",
            "后续将根据真实学生数据改进画像置信度、路径动态调整和资源质量排序。",
        ],
        nums["bullet"],
    )

    doc.add_page_break()
    _add_heading(doc, "附录 A：关键代码定位", 1)
    _add_table(
        doc,
        ["能力", "主要文件"],
        [
            ["学习编排与资产生成", "api/learning_service.py；api/routers/learning.py"],
            ["请求/响应契约", "api/models.py"],
            ["后台生成任务", "commands/learning_commands.py"],
            ["RAG 与语义索引", "forgenote/graphs/*；forgenote/utils/semantic_index.py"],
            [
                "图片与多协议模型",
                "forgenote/ai/image_generation.py；forgenote/ai/model_specs.py",
            ],
            [
                "播客、TTS 与讲解视频",
                "api/podcast_service.py；forgenote/podcasts/robust_creator.py；forgenote/podcasts/video_creator.py",
            ],
            [
                "前端学习工作台",
                "frontend/src/app/(dashboard)/notebooks；frontend/src/components/learning",
            ],
            [
                "专项测试",
                "tests/test_submission_acceptance.py；tests/test_explainer_video.py",
            ],
        ],
        [2800, 6560],
    )
    _add_heading(doc, "附录 B：启动命令", 1)
    for line in (
        "后端：uv run python run_api.py",
        "前端：cd frontend && npm install && npm run dev",
        "Docker：docker compose up -d --build",
        "后端测试：uv run pytest -q",
        "前端测试：cd frontend && npm test -- --reporter=dot",
    ):
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / "开发说明书.docx"
    doc.core_properties.title = "Forgenote 系统开发说明书"
    doc.core_properties.subject = "中国软件杯 A3 赛题配套文档"
    doc.core_properties.author = "Forgenote Team"
    doc.core_properties.keywords = "Forgenote, 多智能体, 多模态, 个性化学习"
    doc.save(output)
    return output


def build_test_manual() -> Path:
    doc = Document()
    nums = _configure_document(doc, "Forgenote · A3 TEST MANUAL")
    _add_cover(
        doc,
        title="Forgenote 测试说明书",
        subtitle="核心验收范围、可复现用例与真实执行证据",
        version="V2.2（字幕与媒体实测版）",
        document_type="测试说明书",
        takeaway=(
            "正文集中说明测试范围、执行环境、8 个官方要求映射用例、讲解视频专项用例、真实结果与适用边界；"
            "附录提供对应测试代码，以及 8/8 和 11/11 两组通过截图。"
        ),
    )

    _add_heading(doc, "1. 测试目的与范围", 1)
    _add_body(
        doc,
        "本次测试验证 Forgenote 的核心学习闭环、带字幕讲解视频生成管线以及相关回归行为。"
        "各项结论均与实际执行命令、用例数量、耗时和证据截图对应。",
    )
    _add_heading(doc, "1.1 本次纳入", 2)
    _add_bullets(
        doc,
        [
            "A3 专项离线验收：画像、多智能体、学习资产、路径、防幻觉、画像更新、流式进度与视频时间轴。",
            "讲解视频专项：脚本视觉提示、真实 TTS 时间点、SRT 字幕、自动换行分段、关键帧图片与真实 MP4 合成。",
            "学习编排回归：LearningService 与 learning API。",
            "后端全量回归：领域、API、数据库适配、命令、工具与 Windows 打包逻辑。",
            "前端单元/组件测试：核心组件、配置、本地化、hooks 与工具函数。",
        ],
        nums["bullet"],
    )
    _add_heading(doc, "1.2 本次不纳入", 2)
    _add_bullets(
        doc,
        [
            "付费大模型、图片和 TTS 供应商的真实网络端到端质量与额度测试。",
            "生产级多租户权限隔离、长期稳定性和大规模并发压力测试。",
            "Windows 安装包在多型号实体设备上的兼容矩阵。",
        ],
        nums["bullet"],
    )

    _add_heading(doc, "2. 测试环境与执行方法", 1)
    _add_table(
        doc,
        ["项目", "本次配置"],
        [
            ["执行日期", "2026-07-19（Asia/Shanghai）"],
            ["操作系统", "Windows"],
            ["后端", "Python 3.11/3.12 兼容环境；pytest；uv"],
            ["前端", "Node.js；Vitest 4.1.8；jsdom"],
            ["专项用例特点", "离线、确定性、无数据库和付费模型凭据依赖"],
            ["证据来源", "pytest/Vitest 实际控制台输出；高清结果截图"],
        ],
        [2600, 6760],
    )
    _add_heading(doc, "2.1 执行命令", 2)
    for line in (
        "uv run pytest -q tests/test_submission_acceptance.py",
        "uv run pytest -q tests/test_explainer_video.py",
        "uv run pytest -q tests/test_learning_service.py tests/test_learning_api.py tests/test_submission_acceptance.py",
        "uv run pytest -q",
        "cd frontend && npm test -- --reporter=dot",
    ):
        p = doc.add_paragraph(style="CodeBlock")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F7F9FC")
        p._p.get_or_add_pPr().append(shd)
        run = p.add_run(line)
        _set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei", size=8.8)

    _add_heading(doc, "3. A3 专项验收用例", 1)
    _add_table(
        doc,
        ["编号", "验证点", "主要断言", "结果"],
        [
            [
                "A3-01",
                "对话式画像",
                "6 个维度；value/evidence 非空；置信度 0–1",
                "PASS",
            ],
            ["A3-02", "多智能体协作", "9 个唯一角色；状态 completed；进度 100", "PASS"],
            [
                "A3-03",
                "多类型资产",
                "7 类输出；Quiz/闪卡/导图/代码/图片契约有效",
                "PASS",
            ],
            ["A3-04", "个性化路径", "4 步有序；每步有目标、活动、资源与检查点", "PASS"],
            [
                "A3-05",
                "防幻觉阻断",
                "来源无正文时 assets=[]，并返回明确阻断提示",
                "PASS",
            ],
            ["A3-06", "画像随学更新", "问答与来源采纳事件进入画像并细化字段", "PASS"],
            ["A3-07", "流式进度", "阶段事件先于最终结果；进度单调不减", "PASS"],
            [
                "A3-08",
                "讲解视频时间轴",
                "关键帧 time_index 来自真实台词 start_time",
                "PASS",
            ],
        ],
        [1000, 2200, 4860, 1300],
        small=True,
    )
    _add_callout(
        doc,
        "设计说明",
        "专项用例直接调用生产编排与媒体合成代码。离线用例不依赖数据库和付费模型；其中一项使用真实音频、真实产品截图和 FFmpeg"
        "生成带中文字幕的可播放 MP4，用于验证本地合成链路。",
        fill=LIGHT_BLUE,
        color=BLUE,
    )

    _add_heading(doc, "4. 执行结果", 1)
    _add_table(
        doc,
        ["测试范围", "结果", "耗时", "备注"],
        [
            ["A3 专项验收", "8 passed", "5.40 s", "1 条依赖弃用 warning"],
            [
                "讲解视频专项",
                "11 passed",
                "35.87 s",
                "含字幕单元测试与 1 个真实 FFmpeg 集成用例",
            ],
            ["学习编排回归", "31 passed", "57.91 s", "2 条依赖弃用 warning"],
            ["后端全量回归", "285 passed", "158.67 s", "2 条依赖弃用 warning"],
            ["前端单元/组件", "10 files / 55 tests passed", "27.10 s", "无失败"],
        ],
        [2200, 2600, 1700, 2860],
        small=True,
    )
    _add_body(
        doc,
        "后端 warning 来自 surreal-commands 的 Pydantic v2 弃用提示，以及 FastAPI TestClient 的 Starlette/httpx 迁移提示；"
        "两者未导致测试失败，依赖升级时应继续跟踪。",
    )

    _add_heading(doc, "5. 结论与限制", 1)
    _add_callout(
        doc,
        "结论",
        "核心学习闭环、带字幕讲解视频管线以及后端和前端回归均无失败。",
        fill="ECFDF5",
        color=GREEN,
    )
    _add_bullets(
        doc,
        [
            "结论仅适用于上述代码版本、环境与测试范围。",
            "真实模型端到端效果应在配置合法 API 凭据后另做演示验收，并记录模型、时间与来源。",
            "多模态质量还需人工检查图片可读性、内容准确性、音频完整性、字幕可读性和视频节奏；自动化测试验证结构与合成契约。",
            "后续建议增加 Docker 冒烟、Windows 安装包自动化、API 性能基线和长任务恢复测试。",
        ],
        nums["bullet"],
    )

    doc.add_page_break()
    _add_heading(doc, "附录 A：A3 专项验收测试源码", 1)
    _add_body(
        doc,
        "源码位置：tests/test_submission_acceptance.py。以下为本次提交对应的完整代码，编号与第三章测试矩阵一致。",
    )
    _add_code_listing(doc, ROOT / "tests" / "test_submission_acceptance.py")

    doc.add_page_break()
    _add_heading(doc, "附录 B：专项验收通过截图", 1)
    _add_body(
        doc,
        "截图保留实际命令、执行时间、8 个用例名称和 tests/failures/errors/skipped/time 字段。",
    )
    _add_image(doc, EVIDENCE, "图 B-1 2026-07-19 A3 专项验收：8/8 通过", width=6.45)

    doc.add_page_break()
    _add_heading(doc, "附录 C：讲解视频专项测试源码", 1)
    _add_body(
        doc,
        "源码位置：tests/test_explainer_video.py。用例覆盖脚本视觉提示、关键帧时间计划、SRT 字幕、自动换行分段、"
        "图片持久化、FFmpeg 字幕烧录、真实 MP4 合成和用户显式选择。",
    )
    _add_code_listing(doc, ROOT / "tests" / "test_explainer_video.py")

    _add_heading(doc, "附录 D：讲解视频专项通过截图", 1)
    _add_image(
        doc, VIDEO_EVIDENCE, "图 D-1 2026-07-19 讲解视频专项：11/11 通过", width=6.45
    )
    _add_image(
        doc,
        ASSETS / "screenshot-explainer-video-details.png",
        "图 D-2 高清界面：已完成单集展示讲解视频标识与真实时间轴关键画面数量",
        width=6.45,
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / "测试说明书.docx"
    doc.core_properties.title = "Forgenote 测试说明书"
    doc.core_properties.subject = "中国软件杯 A3 赛题测试证据"
    doc.core_properties.author = "Forgenote Team"
    doc.core_properties.keywords = "Forgenote, 测试, pytest, Vitest, A3"
    doc.save(output)
    return output


if __name__ == "__main__":
    dev = build_development_manual()
    test = build_test_manual()
    print(dev)
    print(test)
